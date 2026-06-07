"""
bulletin.py — Turn the weekly .docx into clean, structured events.
==================================================================

Flow:
  1. Pull plain text out of the .docx (python-docx) — paragraphs AND tables.
  2. Hand it to sanitize.py so student data is stripped.
  3. Send ONLY the sanitized text to the API and ask for strict JSON events.

The API is good at the messy-human part ("WOW Celebrations Wednesday in MU Room
TK-4th @ 8:30am 5th-8th @ 9:30am") that regex is bad at. We give it the week the
bulletin covers so it can resolve "Wednesday" into a real date.

Everything here fails soft: no key / no network / bad response → empty event
list plus a warning, never a crash.
"""

from __future__ import annotations

import datetime as dt
import json
import re
from dataclasses import dataclass, field

import config
import sanitize


@dataclass
class BulletinResult:
    events: list[dict]            # normalized event dicts (see schema below)
    week_of: str | None           # ISO date string for the Monday of the week
    sanitize_summary: str         # PII-free description of what was stripped
    warning: str | None = None
    announcements: list[dict] = field(default_factory=list)  # ticker notices: [{"text": ...}]


# The event schema we ask the model to fill. Keeping it here documents the
# contract the display depends on:
#   title     short label              ("WOW Celebration TK–4")
#   date      "YYYY-MM-DD" or null      (null when truly undated)
#   day       weekday name or null      ("Wednesday")
#   start     "HH:MM" 24h or null
#   end       "HH:MM" 24h or null
#   location  room/place or null        ("MU Room")
#   category  one of CATEGORIES         (drives icon/color on the board)
#   audience  who it's for or null      ("TK-4th", "Classified staff")
CATEGORIES = ["celebration", "meeting", "deadline", "trip", "food", "sports", "general"]


def extract_text(source) -> str:
    """
    Read a .docx into plain text, including tables (the yard-duty grid is a
    table, and so are some announcements). python-docx exposes paragraphs and
    tables separately, so we walk both. `source` may be a file path or an
    in-memory binary stream (e.g. a BytesIO of a downloaded file)."""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError("python-docx not installed — run pip install python-docx") from exc

    document = docx.Document(source)
    chunks: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                chunks.append(line)

    return "\n".join(chunks)


# ---- Google Doc source -----------------------------------------------------
# Instead of a downloaded .docx, the bulletin can be pulled straight from a
# Google Doc link. We export it to plain text via the Drive API using the SAME
# service account that reads the yard-duty Sheet — so the doc stays private
# (shared only with that service account), never made link-public. The bulletin
# still contains student data, so this MUST stay a private, authenticated fetch.
_GDOC_ID_RE = re.compile(r"/document/d/([a-zA-Z0-9_-]+)")
_BARE_ID_RE = re.compile(r"^[a-zA-Z0-9_-]{20,}$")


def _extract_doc_id(url: str) -> str | None:
    url = (url or "").strip()
    m = _GDOC_ID_RE.search(url)
    if m:
        return m.group(1)
    return url if _BARE_ID_RE.match(url) else None  # they may have pasted a bare ID


def _service_account_email() -> str:
    try:
        with open(config.GOOGLE_SERVICE_ACCOUNT_FILE) as fh:
            return json.load(fh).get("client_email", "the service account")
    except Exception:
        return "the service account"


def _drive_session():
    """An authorized requests session for the Drive API, reusing the Sheet's
    service account (read-only scope)."""
    from google.oauth2 import service_account
    from google.auth.transport.requests import AuthorizedSession

    creds = service_account.Credentials.from_service_account_file(
        config.GOOGLE_SERVICE_ACCOUNT_FILE,
        scopes=["https://www.googleapis.com/auth/drive.readonly"],
    )
    return AuthorizedSession(creds)


def _check_drive_response(resp) -> None:
    """Turn Drive's auth/permission errors into messages that say what to fix."""
    if resp.status_code == 404:
        raise RuntimeError("Not found — double-check the link or folder ID.")
    if resp.status_code == 403:
        raise RuntimeError(
            f"Access denied. Share the doc/folder with {_service_account_email()} "
            "as Viewer, and make sure the Drive API is enabled for the project."
        )
    resp.raise_for_status()


def export_doc_text(doc_id: str) -> str:
    """Export a Google Doc (by ID) to plain text. supportsAllDrives so it works
    for files living in a Shared Drive."""
    session = _drive_session()
    resp = session.get(
        f"https://www.googleapis.com/drive/v3/files/{doc_id}/export",
        params={"mimeType": "text/plain", "supportsAllDrives": "true"}, timeout=30,
    )
    _check_drive_response(resp)
    resp.encoding = "utf-8"
    return resp.text


# Bulletin files come in two shapes: a native Google Doc, or an uploaded Word
# file. Modern Drive does NOT auto-convert uploads, so a .docx stays a .docx —
# we have to look for, and read, both.
GDOC_MIME = "application/vnd.google-apps.document"
WORD_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def list_drive_folder_docs(folder_id: str) -> list[dict]:
    """List the bulletin files in a folder (Shared Drive aware), newest first.
    Returns dicts with id, name, modifiedTime, mimeType. Includes BOTH native
    Google Docs and uploaded Word (.docx) files, since either may be dropped in.
    corpora=allDrives makes the folder search reach into Shared Drives too."""
    session = _drive_session()
    resp = session.get(
        "https://www.googleapis.com/drive/v3/files",
        params={
            "q": (f"'{folder_id}' in parents and trashed=false "
                  f"and (mimeType='{GDOC_MIME}' or mimeType='{WORD_MIME}')"),
            "fields": "files(id,name,modifiedTime,mimeType)",
            "orderBy": "modifiedTime desc",
            "pageSize": "100",
            "corpora": "allDrives",
            "supportsAllDrives": "true",
            "includeItemsFromAllDrives": "true",
        },
        timeout=30,
    )
    _check_drive_response(resp)
    return resp.json().get("files", [])


def fetch_drive_file_text(file: dict) -> str:
    """Plain text from a Drive file dict (needs id + mimeType): export a Google
    Doc, or download + extract an uploaded Word file."""
    mime = file.get("mimeType", "")
    if mime == GDOC_MIME:
        return export_doc_text(file["id"])
    if mime == WORD_MIME:
        return _download_docx_text(file["id"])
    raise RuntimeError(f"Unsupported bulletin file type ({mime or 'unknown'}).")


def _download_docx_text(file_id: str) -> str:
    """Download an uploaded .docx (alt=media) and extract its text in-memory."""
    import io
    session = _drive_session()
    resp = session.get(
        f"https://www.googleapis.com/drive/v3/files/{file_id}",
        params={"alt": "media", "supportsAllDrives": "true"}, timeout=30,
    )
    _check_drive_response(resp)
    return extract_text(io.BytesIO(resp.content))


def fetch_google_doc(url: str) -> tuple[str, str]:
    """Export a Google Doc to plain text via the Drive API. Returns (title, text).
    The doc (or a folder it lives in) must be shared with the service account as
    Viewer, and the Drive API must be enabled for the project."""
    doc_id = _extract_doc_id(url)
    if not doc_id:
        raise RuntimeError("Couldn't find a Google Doc ID in that link — paste the full /document/d/… URL.")

    session = _drive_session()
    meta = session.get(
        f"https://www.googleapis.com/drive/v3/files/{doc_id}",
        params={"fields": "name,mimeType", "supportsAllDrives": "true"}, timeout=30,
    )
    _check_drive_response(meta)
    info = meta.json()
    if info.get("mimeType") != "application/vnd.google-apps.document":
        raise RuntimeError("That link isn't a Google Doc — only Docs export to text this way.")

    return info.get("name", ""), export_doc_text(doc_id)


_MONTHS = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}


def _week_from_string(s: str | None, numeric_ok: bool = False) -> dt.date | None:
    """Pull a date out of a string and return it (current year if no year given),
    or None. Always recognizes a written month + day ('June 1', 'Week of Jun 1',
    'May 4-8'). With numeric_ok=True it ALSO accepts 'M/D', 'M-D', 'M/D/YY',
    'M/D/YYYY' — use that for filenames/Doc titles, but NOT for scanning body
    text, where '3/4' is more likely a fraction than a date."""
    text = s or ""
    m = re.search(
        r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+(\d{1,2})",
        text, re.IGNORECASE,
    )
    if m:
        try:
            return dt.date(dt.date.today().year, _MONTHS[m.group(1).lower()[:3]], int(m.group(2)))
        except ValueError:
            return None

    if numeric_ok:
        m = re.search(r"\b(\d{1,2})[/-](\d{1,2})(?:[/-](\d{2,4}))?\b", text)
        if m:
            month, day = int(m.group(1)), int(m.group(2))
            year = int(m.group(3)) if m.group(3) else dt.date.today().year
            if year < 100:
                year += 2000
            try:
                return dt.date(year, month, day)
            except ValueError:
                return None
    return None


def week_of_from_filename(docx_path: str, fallback: dt.date | None = None) -> dt.date:
    """The files are named like "Week of June 1.docx", so the date is right
    there. Fall back to (this) Monday if the name doesn't cooperate."""
    name = re.sub(r"\.docx$", "", docx_path.split("/")[-1], flags=re.IGNORECASE)
    return _week_from_string(name, numeric_ok=True) or fallback or _monday_of(dt.date.today())


def _monday_of(d: dt.date) -> dt.date:
    return d - dt.timedelta(days=d.weekday())


def _build_prompt(safe_text: str, week_start: dt.date) -> str:
    """The extraction instructions. We pin the week so days resolve to dates."""
    week_end = week_start + dt.timedelta(days=6)
    return f"""You are reading a school staff bulletin to populate a lobby display.
The bulletin covers the week of {week_start:%A, %B %-d, %Y} through
{week_end:%A, %B %-d, %Y}.

Return ONLY a JSON object — no prose, no markdown fences:
{{
  "events": [ /* dated/scheduled happenings, schema below */ ],
  "announcements": [ /* short general notices for a scrolling ticker, strings */ ]
}}

Each EVENT object:
{{
  "title": "short human label, max ~6 words",
  "date": "YYYY-MM-DD or null",
  "day": "weekday name or null",
  "start": "HH:MM 24-hour or null",
  "end": "HH:MM 24-hour or null",
  "location": "room or place, or null",
  "category": one of {CATEGORIES},
  "audience": "who it's for, or null"
}}

Event rules:
- Resolve weekday names to real dates using the week above. Dates with an
  explicit month/day (e.g. "June 9") may fall outside that week — keep them.
- Split one line into multiple events when it clearly lists several (e.g. an
  event with different times for different grades → one event per grade).
- Skip pure instructions, policies, and material-prep notes — only real
  dated/scheduled happenings.
- Skip anything that is just a hyperlink or a form.
- If a time is a range like "11:00-12:30", fill both start and end.
- Be conservative with category; use "general" when unsure.

ANNOUNCEMENTS are the general notices, reminders, and messages to staff that
are NOT a specific dated event — the kind of thing that scrolls along a ticker.
- Each item is a plain string, max ~12 words, lightly cleaned up for display.
- Examples: "Picture day forms are due Friday", "Please lock classroom doors at dismissal".
- Do NOT include anyone's name, student information, links, or forms.
- Do NOT repeat something already captured as an event.
- If there are none, return an empty array.

BULLETIN TEXT:
\"\"\"
{safe_text}
\"\"\"
"""


def _coerce_payload(raw: str) -> tuple[list[dict], list[dict]]:
    """Parse the model's reply into (events, announcements), tolerating stray
    fences and an older array-only shape (treated as events with no notices)."""
    cleaned = raw.strip()
    # Strip ```json … ``` if the model added them despite instructions.
    cleaned = re.sub(r"^```[a-zA-Z]*\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    data = json.loads(cleaned)
    if isinstance(data, list):           # legacy: a bare array of events
        raw_events, raw_notices = data, []
    elif isinstance(data, dict):
        raw_events = data.get("events", [])
        raw_notices = data.get("announcements", [])
    else:
        raise ValueError("expected a JSON object with events/announcements")

    # Light validation/normalization so the display never sees surprises.
    events: list[dict] = []
    for item in raw_events:
        if not isinstance(item, dict) or not item.get("title"):
            continue
        category = item.get("category")
        if category not in CATEGORIES:
            category = "general"
        events.append(
            {
                "title": str(item["title"]).strip(),
                "date": item.get("date"),
                "day": item.get("day"),
                "start": item.get("start"),
                "end": item.get("end"),
                "location": item.get("location"),
                "category": category,
                "audience": item.get("audience"),
            }
        )

    # Notices come back as plain strings; wrap them in the {"text": ...} shape
    # the ticker reads, dropping blanks.
    announcements: list[dict] = []
    for note in raw_notices:
        text = (note if isinstance(note, str) else note.get("text", "")).strip()
        if text:
            announcements.append({"text": text})

    return events, announcements


def parse(raw_text: str, week_hint: str | None = None) -> BulletinResult:
    """Shared pipeline: derive week → sanitize → (confirm) → API → events/notices.

    `raw_text` is the bulletin's plain text, from either a .docx (extract_text)
    or a Google Doc (fetch_google_doc). `week_hint` is the filename or doc title
    — that's where "Week of <date>" usually lives; we fall back to scanning the
    text, then to this Monday."""
    week_start = (_week_from_string(week_hint)
                  or _week_from_string((raw_text or "")[:600])
                  or _monday_of(dt.date.today()))

    safe_text, report = sanitize.sanitize_bulletin_text(raw_text)
    summary = report.summary()

    # The human confirmation gate — the real guarantee for interactive runs.
    if config.REQUIRE_SEND_CONFIRMATION:
        if not _confirm_send(safe_text):
            return BulletinResult([], week_start.isoformat(), summary,
                                  "send not confirmed — skipped API event extraction")

    events, announcements, warning = extract_events_notices(safe_text, week_start)
    return BulletinResult(events, week_start.isoformat(), summary, warning, announcements)


def extract_events_notices(
    safe_text: str, week_start: dt.date
) -> tuple[list[dict], list[dict], str | None]:
    """API call → (events, announcements, warning). Assumes safe_text is ALREADY
    sanitized — the autonomous job sanitizes and runs its fail-safe first, then
    calls this. Returns empties + a warning string on any failure instead of
    raising, so the caller can choose NOT to blank the board on a hiccup."""
    if not config.ANTHROPIC_API_KEY:
        return [], [], "no ANTHROPIC_API_KEY — skipped event extraction"
    try:
        import anthropic
    except ImportError:
        return [], [], "anthropic SDK not installed — skipped event extraction"
    try:
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        message = client.messages.create(
            model=config.ANTHROPIC_MODEL,
            max_tokens=8000,
            messages=[{"role": "user", "content": _build_prompt(safe_text, week_start)}],
        )
        reply = "".join(block.text for block in message.content if block.type == "text")
        events, announcements = _coerce_payload(reply)
        return events, announcements, None
    except Exception as exc:
        return [], [], f"event extraction failed ({exc})"


def _confirm_send(safe_text: str) -> bool:
    """
    Show EXACTLY what will be sent to the API and require an explicit yes.
    Also surface any lingering roster-pattern warnings in the clear.
    """
    print("\n" + "=" * 70)
    print("ABOUT TO SEND THIS TEXT TO THE API. Student data should NOT appear.")
    print("=" * 70)
    print(safe_text if safe_text else "(empty)")
    print("=" * 70)

    warnings = sanitize.looks_clean(safe_text)
    if warnings:
        print("\n⚠️  HEADS UP — these lines still look roster-ish:")
        for w in warnings:
            print(f"   • {w}")
        print("   Review above. If you see student names, type anything but 'yes'.\n")

    answer = input('Send this to the API? Type "yes" to proceed: ').strip().lower()
    return answer == "yes"
